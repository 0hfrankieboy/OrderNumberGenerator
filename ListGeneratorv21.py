import sys
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document
from docx.shared import Inches

'''
    This is a program for generating order numbers for Prime Time International and inserting them into a word document.
    It takes a beginning number and the amount that you would like to create.
    Written by Frank Caniza Sept 2016.
                                        '''
def main():    
    first = firstNumber()
    amount = amountGenerated()
    end = first + amount
    myList = [first]
            
    while True:
        if first != end:
            myList.append(first + 1)
            first = first + 1
        elif first == end:
            print(myList)
            makeFile(myList)
            break

def firstNumber():
    while True:
        try:
            first = int(input("Beginning number for sequence: "))
        except ValueError:
            print("Please enter a number")
            continue
        else:        
            break
    return first
            
def amountGenerated():
    while True:
        try:    
            amount = int(input("How many order numbers would you like to generate?: "))
        except ValueError:
            print("Please enter a number")
            continue
        if amount > 1000:
            print("Please do not exceed 1000 order numbers")
            continue
        else:
            confirm = input("Are you sure you want to print this many numbers? Y or N: ")
            if confirm.lower() == "n":
                continue
            if confirm.lower() == "y":
                break
    return amount
    
def makeFile(myList): #creates and formats the list into word file      
    #initialize document
    doc = docx.Document()
    #get input for document
    salesperson = doc.add_paragraph(input("Enter salespersons name: "))
    docName = input("Enter file name: ") + ".docx"
    #returns the list without brackets and commas
    myList = "\n".join(map(str, myList)) 
    #add salespersons name
    salesperson.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    #add image parameters
    doc.add_picture('ptilogo.png', width=Inches(1))
    logoParagraph = doc.paragraphs[-1] # assigns position of picture
    logoParagraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #add heading parameters
    doc.add_heading("PRIME TIME ORDER NUMBERS", 0)
    headingParagraph = doc.paragraphs[-1] #assigns position of header
    headingParagraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #add list to document
    list = doc.add_paragraph(myList)
    list.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.save(docName)

main()



     
        
        
