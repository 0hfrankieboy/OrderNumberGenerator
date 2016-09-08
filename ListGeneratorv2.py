import sys
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document
from docx.shared import Inches

'''
    This is a program for generating order numbers for Prime Time International and inserting them into a word document.
    It takes a beginning number and the amount that you would like to create.
    Written by Frank Caniza Sept 2016.
                                        '''



def defineAmount():
    amount = int(input("How many order numbers would you like to generate?: "))
    if amount > 100:
        ConfirmCounter = 0
        while ConfirmCounter < 1:
            confirm = input("Are you sure you want to print this many numbers? Y or N: ")
            if confirm.lower() == "y":
                ConfirmCounter = 3
            if confirm.lower() == "n":
               amount = int(input("How many order numbers would you like to generate?: "))
               ConfirmCounter += 1
    return amount

def makeFile(myList): #Creates and formats the list into word file      
    myList = "\n".join(map(str, myList)) #returns the list without brackets and commas
    docName = input("Enter file name: ")
    docName = docName + '.docx'
    #initialize document
    doc = docx.Document()
    #add image parameters
    doc.add_picture('ptilogo.png', width=Inches(1))
    logoParagraph = doc.paragraphs[-1] # assigns position of picture
    logoParagraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #add heading parameters
    doc.add_heading("PRIME TIME ORDER NUMBERS", 0)
    headingParagraph = doc.paragraphs[-1] #assigns position of header
    #add list to document
    headingParagraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    list = doc.add_paragraph(myList)
    list.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.save(docName)
    
        
first = int(input("Beginning number for sequence: "))
end = first + defineAmount()
myList = [first]
            
while True:
    if first != end:
        myList.append(first + 1)
        first = first + 1
    elif first == end:
        print(myList)
        makeFile(myList)
        break